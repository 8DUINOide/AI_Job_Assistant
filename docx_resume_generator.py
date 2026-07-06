"""
DOCX Resume Generator — Uses the developer-resume-template.docx as the base 
template and replaces all content with the user's profile data while preserving
the exact formatting (fonts, sizes, bold, tab stops, section dividers).
"""
import io
import os
import copy
from docx import Document
from docx.shared import Pt, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Path to the template file
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                             'android', 'developer-resume-template.docx')

# Font sizes matching the template
NAME_SIZE = Pt(16)
SECTION_HEADER_SIZE = Pt(11)
BODY_SIZE = Pt(10.5)


def _clear_paragraphs(doc):
    """Remove all existing paragraphs from the document body."""
    body = doc.element.body
    for p in list(body.findall(qn('w:p'))):
        body.remove(p)


def _add_paragraph(doc, text="", bold=False, size=BODY_SIZE, alignment=None, space_after=Pt(0), space_before=Pt(0)):
    """Add a paragraph with specific formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = Pt(13)
    
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = size
        run.font.name = 'Calibri'
    
    return p


def _add_two_column_line(doc, left_text, right_text, left_bold=False, right_bold=False, 
                         size=BODY_SIZE, space_after=Pt(0), space_before=Pt(0)):
    """Add a line with left-aligned and right-aligned text using a tab stop."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = Pt(13)
    
    # Add a right-aligned tab stop at the right margin
    from docx.oxml import OxmlElement
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '10080')  # ~7 inches in twips (right margin)
    tabs.append(tab)
    pPr.append(tabs)
    
    # Left text
    run_left = p.add_run(left_text)
    run_left.bold = left_bold
    run_left.font.size = size
    run_left.font.name = 'Calibri'
    
    # Tab separator
    run_tab = p.add_run('\t')
    run_tab.font.size = size
    
    # Right text
    run_right = p.add_run(right_text)
    run_right.bold = right_bold
    run_right.font.size = size
    run_right.font.name = 'Calibri'
    
    return p


def _add_section_header(doc, title):
    """Add a section header with a horizontal line (matching template style)."""
    p = _add_paragraph(doc, title, bold=True, size=SECTION_HEADER_SIZE, 
                       space_before=Pt(8), space_after=Pt(2))
    
    # Add bottom border to simulate the horizontal rule
    from docx.oxml import OxmlElement
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    return p


def _add_bullet_point(doc, text, size=BODY_SIZE):
    """Add a bullet point."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(1)
    pf.space_before = Pt(0)
    pf.line_spacing = Pt(13)
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.15)
    
    run = p.add_run('● ')
    run.font.size = size
    run.font.name = 'Calibri'
    
    run2 = p.add_run(text)
    run2.font.size = size
    run2.font.name = 'Calibri'
    
    return p


def generate_resume_docx(profile_data):
    """
    Generate a DOCX resume from the template using the user's profile data.
    
    Args:
        profile_data: dict with keys matching the template sections:
            - personal_info: {name, location, phone, email, linkedin, github}
            - education: [{university, location, degree, year, details, coursework}]
            - experience: [{company, company_note, location, title, date_range, bullets}]
            - projects: [{name, year, bullets}]
            - activities: [{organization, location, role, date_range, bullets}]
            - additional: {technical_skills, certifications, languages}
    
    Returns:
        BytesIO buffer containing the generated DOCX file
    """
    # Load the template
    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
    
    # Clear existing content
    _clear_paragraphs(doc)
    
    pi = profile_data.get('personal_info', {})
    
    # === HEADER: Name ===
    name = pi.get('name', '')
    _add_paragraph(doc, name.upper(), bold=True, size=NAME_SIZE, 
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    
    # === HEADER: Contact line ===
    contact_parts = []
    if pi.get('location'):
        contact_parts.append(pi['location'])
    if pi.get('phone'):
        contact_parts.append(pi['phone'])
    if pi.get('email'):
        contact_parts.append(pi['email'])
    if pi.get('linkedin'):
        contact_parts.append(pi['linkedin'])
    if pi.get('github'):
        contact_parts.append(pi['github'])
    
    contact_line = ' | '.join(contact_parts)
    _add_paragraph(doc, contact_line, bold=False, size=BODY_SIZE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    
    # === EDUCATION ===
    education = profile_data.get('education', [])
    if education:
        _add_section_header(doc, 'EDUCATION')
        
        for edu in education:
            # University + Location
            _add_two_column_line(doc, 
                                edu.get('university', '').upper(), 
                                edu.get('location', ''),
                                left_bold=True, space_before=Pt(4))
            
            # Degree + Year
            _add_two_column_line(doc,
                                edu.get('degree', ''),
                                edu.get('year', ''),
                                space_after=Pt(1))
            
            # Details (major/minor, GWA, etc.)
            if edu.get('details'):
                _add_paragraph(doc, edu['details'], size=BODY_SIZE, space_after=Pt(1))
            
            # Coursework
            if edu.get('coursework'):
                _add_paragraph(doc, f"Relevant Coursework: {edu['coursework']}", 
                              size=BODY_SIZE, space_after=Pt(2))
    
    # === WORK EXPERIENCE ===
    experience = profile_data.get('experience', [])
    if experience:
        _add_section_header(doc, 'WORK EXPERIENCE')
        
        for exp in experience:
            # Company + Location
            company_text = exp.get('company', '').upper()
            if exp.get('company_note'):
                company_text += f" ({exp['company_note']})"
            
            _add_two_column_line(doc,
                                company_text,
                                exp.get('location', ''),
                                left_bold=True, space_before=Pt(4))
            
            # Title + Date
            _add_two_column_line(doc,
                                exp.get('title', ''),
                                exp.get('date_range', ''),
                                space_after=Pt(2))
            
            # Bullets
            for bullet in exp.get('bullets', []):
                if bullet.strip():
                    _add_bullet_point(doc, bullet.strip())
    
    # === PROJECTS ===
    projects = profile_data.get('projects', [])
    if projects:
        _add_section_header(doc, 'PROJECTS')
        
        for proj in projects:
            # Project name + Year
            _add_two_column_line(doc,
                                proj.get('name', '').upper(),
                                proj.get('year', ''),
                                left_bold=True, space_before=Pt(4))
            
            # Bullets
            for bullet in proj.get('bullets', []):
                if bullet.strip():
                    _add_bullet_point(doc, bullet.strip())
    
    # === ACTIVITIES ===
    activities = profile_data.get('activities', [])
    if activities:
        _add_section_header(doc, 'ACTIVITIES')
        
        for act in activities:
            # Organization + Location
            _add_two_column_line(doc,
                                act.get('organization', '').upper(),
                                act.get('location', ''),
                                left_bold=True, space_before=Pt(4))
            
            # Role + Date range
            _add_two_column_line(doc,
                                act.get('role', ''),
                                act.get('date_range', ''),
                                space_after=Pt(2))
            
            # Bullets
            for bullet in act.get('bullets', []):
                if bullet.strip():
                    _add_bullet_point(doc, bullet.strip())
    
    # === ADDITIONAL ===
    additional = profile_data.get('additional', {})
    has_additional = (additional.get('technical_skills') or 
                     additional.get('certifications') or 
                     additional.get('languages'))
    
    if has_additional:
        _add_section_header(doc, 'ADDITIONAL')
        
        if additional.get('technical_skills'):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(2)
            pf.line_spacing = Pt(13)
            run_label = p.add_run('Technical Skills: ')
            run_label.bold = True
            run_label.font.size = BODY_SIZE
            run_label.font.name = 'Calibri'
            run_val = p.add_run(additional['technical_skills'])
            run_val.font.size = BODY_SIZE
            run_val.font.name = 'Calibri'
        
        if additional.get('certifications'):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(2)
            pf.line_spacing = Pt(13)
            run_label = p.add_run('Certifications: ')
            run_label.bold = True
            run_label.font.size = BODY_SIZE
            run_label.font.name = 'Calibri'
            run_val = p.add_run(additional['certifications'])
            run_val.font.size = BODY_SIZE
            run_val.font.name = 'Calibri'
        
        if additional.get('languages'):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(2)
            pf.line_spacing = Pt(13)
            run_label = p.add_run('Languages: ')
            run_label.bold = True
            run_label.font.size = BODY_SIZE
            run_label.font.name = 'Calibri'
            run_val = p.add_run(additional['languages'])
            run_val.font.size = BODY_SIZE
            run_val.font.name = 'Calibri'
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def convert_user_profile_to_template_data(profile):
    """
    Converts a UserProfile (from Firestore/Android) into the flat structure 
    expected by generate_resume_docx.
    """
    pi = profile.get('personalInfo') or profile.get('personal_info', {})
    first_name = pi.get('firstName') or pi.get('first_name', '')
    last_name = pi.get('lastName') or pi.get('last_name', '')
    
    template_data = {
        'personal_info': {
            'name': f"{first_name} {last_name}".strip(),
            'location': pi.get('location', ''),
            'phone': pi.get('phone', ''),
            'email': pi.get('email', ''),
            'linkedin': pi.get('linkedinUrl') or pi.get('linkedin_url', ''),
            'github': pi.get('portfolioUrl') or pi.get('portfolio_url', '')
        },
        'education': [],
        'experience': [],
        'projects': [],
        'activities': profile.get('activities', []),
        'additional': {
            'technical_skills': ', '.join(profile.get('skills', [])),
            'certifications': ', '.join(profile.get('certifications', [])),
            'languages': ''
        }
    }
    
    # Map education
    for edu in profile.get('education', []):
        template_data['education'].append({
            'university': edu.get('university', ''),
            'location': '',
            'degree': edu.get('degree', ''),
            'year': edu.get('graduationYear') or edu.get('graduation_year', ''),
            'details': '',
            'coursework': ''
        })
    
    # Map experience
    for exp in profile.get('experience', []):
        start = exp.get('startDate') or exp.get('start_date', '')
        end = exp.get('endDate') or exp.get('end_date', '')
        date_range = f"{start} – {end}" if start and end else (start or end)
        
        desc = exp.get('description', '')
        bullets = [b.strip() for b in desc.split('. ') if b.strip()]
        
        template_data['experience'].append({
            'company': exp.get('company', ''),
            'company_note': '',
            'location': exp.get('location', ''),
            'title': exp.get('title', ''),
            'date_range': date_range,
            'bullets': bullets
        })
    
    # Map projects
    for proj in profile.get('projects', []):
        desc = proj.get('description', '')
        bullets = [b.strip() for b in desc.split('. ') if b.strip()]
        
        template_data['projects'].append({
            'name': proj.get('title', ''),
            'year': '',
            'bullets': bullets
        })
    
    return template_data
